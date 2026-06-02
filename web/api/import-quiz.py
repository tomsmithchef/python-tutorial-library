import io
import json
import os
import re
from email import policy
from email.parser import BytesParser
from http.server import BaseHTTPRequestHandler


MAX_UPLOAD_BYTES = 8 * 1024 * 1024


def json_response(handler, status, payload):
    body = json.dumps(payload).encode("utf-8")
    handler.send_response(status)
    handler.send_header("Content-Type", "application/json; charset=utf-8")
    handler.send_header("Content-Length", str(len(body)))
    handler.end_headers()
    handler.wfile.write(body)


def extract_upload(headers, body):
    content_type = headers.get("Content-Type", "")
    if "multipart/form-data" not in content_type:
        raise ValueError("Upload a DOCX or PDF file using multipart/form-data.")

    message_bytes = (
        f"Content-Type: {content_type}\r\n"
        f"Content-Length: {len(body)}\r\n\r\n"
    ).encode("utf-8") + body
    message = BytesParser(policy=policy.default).parsebytes(message_bytes)

    for part in message.iter_parts():
        filename = part.get_filename()
        if filename:
            payload = part.get_payload(decode=True) or b""
            return filename, payload

    raise ValueError("No uploaded file was found.")


def extract_docx_text(file_bytes):
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("DOCX import dependency is not installed.") from exc

    document = Document(io.BytesIO(file_bytes))
    lines = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" ".join(cells))
    return "\n".join(lines)


def extract_pdf_text(file_bytes):
    try:
        import fitz
    except Exception as exc:
        raise RuntimeError("PDF import dependency is not installed.") from exc

    text_parts = []
    with fitz.open(stream=file_bytes, filetype="pdf") as document:
        for page in document:
            page_text = page.get_text("text").strip()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def split_answer_section(text):
    match = re.search(r"(?im)^\s*(answers?|answer\s+key)\s*:?\s*$", text)
    if not match:
        return text, ""
    return text[:match.start()], text[match.end():]


def clean_text(value):
    return re.sub(r"\s+", " ", str(value or "")).strip()


def strip_question_number(line):
    return re.sub(r"^\s*\d+\s*[\.\)]?\s*", "", line).strip()


def extract_answers(answer_text):
    answers = {}
    lines = [line.strip() for line in answer_text.splitlines() if line.strip()]

    table_tokens = []
    for line in lines:
        cleaned = clean_text(line).strip()
        if re.match(r"^(q#|question|answer|answers?)$", cleaned, re.IGNORECASE):
            continue
        if re.match(r"^\d+$", cleaned) or re.match(r"^[A-E](?:\s*,\s*[A-E])*$", cleaned, re.IGNORECASE):
            table_tokens.append(cleaned.upper())

    index = 0
    while index < len(table_tokens) - 1:
        qnum_token = table_tokens[index]
        answer_token = table_tokens[index + 1]
        if qnum_token.isdigit() and re.match(r"^[A-E](?:\s*,\s*[A-E])*$", answer_token, re.IGNORECASE):
            answers[int(qnum_token)] = {
                "labels": [label.upper() for label in re.findall(r"[A-E]", answer_token, re.IGNORECASE)],
                "written": [],
                "explanation": ""
            }
            index += 2
        else:
            index += 1

    if answers:
        return answers

    sequential = 1

    for line in lines:
        numbered = re.match(r"^(\d+)\s*[\.\)\:-]?\s*(.+)$", line)
        if numbered:
            qnum = int(numbered.group(1))
            payload = numbered.group(2).strip()
        else:
            qnum = sequential
            payload = line
            sequential += 1

        letters = re.match(r"^([A-E](?:\s*,\s*[A-E])*)\.?\s*(.*)$", payload, re.IGNORECASE)
        if letters:
            labels = [label.upper() for label in re.findall(r"[A-E]", letters.group(1), re.IGNORECASE)]
            explanation = clean_text(letters.group(2))
            answers[qnum] = {
                "labels": labels,
                "written": [],
                "explanation": explanation or ""
            }
            continue

        written_parts = [part.strip() for part in re.split(r"[;|]", payload) if part.strip()]
        if written_parts:
            answers[qnum] = {
                "labels": [],
                "written": written_parts,
                "explanation": ""
            }

    return answers


def looks_like_question_start(line):
    return bool(re.match(r"^\s*\d+(?:\s*[\.\)]\s+|\s+).+", line))


def split_question_blocks(question_text):
    lines = [line.rstrip() for line in question_text.splitlines()]
    starts = []
    for index, line in enumerate(lines):
        stripped = line.strip()
        if looks_like_question_start(stripped):
            starts.append(index)

    blocks = []
    for offset, start in enumerate(starts):
        end = starts[offset + 1] if offset + 1 < len(starts) else len(lines)
        block = "\n".join(lines[start:end]).strip()
        qnum_match = re.match(r"^\s*(\d+)", block)
        if block and qnum_match:
            blocks.append((int(qnum_match.group(1)), block))
    return blocks


def parse_options(lines):
    options = []
    question_lines = []
    current_label = None
    current_text = []

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue

        option_match = re.match(r"^([A-E])[\.\)]\s*(.+)$", line, re.IGNORECASE)
        if option_match:
            if current_label and current_text:
                options.append({
                    "label": current_label,
                    "text": clean_text(" ".join(current_text))
                })
            current_label = option_match.group(1).upper()
            current_text = [option_match.group(2).strip()]
            continue

        if current_label:
            current_text.append(line)
        else:
            question_lines.append(line)

    if current_label and current_text:
        options.append({
            "label": current_label,
            "text": clean_text(" ".join(current_text))
        })

    return clean_text(" ".join(question_lines)), options


def parse_questions(raw_text, title):
    text = raw_text.replace("\r\n", "\n").replace("\r", "\n")
    question_text, answer_text = split_answer_section(text)
    answers = extract_answers(answer_text)
    if not answers:
        raise ValueError("This quiz document needs an ANSWERS section with an answer key.")

    blocks = split_question_blocks(question_text)
    if not blocks:
        raise ValueError("No numbered questions were found in the uploaded document.")

    questions = []
    warnings = []
    for qnum, block in blocks:
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue
        lines[0] = strip_question_number(lines[0])
        question, options = parse_options(lines)
        answer = answers.get(qnum)
        if not answer:
            warnings.append(f"Question {qnum} did not have a matching answer key entry.")
            continue

        if options and answer["labels"]:
            valid_labels = {option["label"] for option in options}
            correct = [label for label in answer["labels"] if label in valid_labels]
            if not correct:
                warnings.append(f"Question {qnum} answer labels did not match its options.")
                continue
            questions.append({
                "id": f"import-q{qnum}",
                "type": "multi" if len(correct) > 1 else "choice",
                "text": question,
                "options": options,
                "correctLabels": correct,
                "acceptedAnswers": [],
                "hint": "",
                "explanation": answer.get("explanation", "")
            })
        elif answer["written"]:
            questions.append({
                "id": f"import-q{qnum}",
                "type": "written",
                "text": question,
                "options": [],
                "correctLabels": [],
                "acceptedAnswers": answer["written"],
                "hint": "",
                "explanation": answer.get("explanation", "")
            })

    if not questions:
        raise ValueError("No usable questions were found after matching questions to the answer key.")

    return {
        "title": title,
        "questions": questions,
        "warnings": warnings
    }


class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            content_length = int(self.headers.get("Content-Length", "0"))
            if content_length <= 0:
                json_response(self, 400, {"error": "Upload a DOCX or PDF quiz file."})
                return
            if content_length > MAX_UPLOAD_BYTES:
                json_response(self, 413, {"error": "The uploaded file is too large. Keep quiz files under 8 MB."})
                return

            body = self.rfile.read(content_length)
            filename, file_bytes = extract_upload(self.headers, body)
            ext = os.path.splitext(filename)[1].lower()
            if ext not in {".docx", ".pdf"}:
                json_response(self, 400, {"error": "Unsupported format. Upload a .docx or .pdf quiz file."})
                return

            if ext == ".docx":
                raw_text = extract_docx_text(file_bytes)
            else:
                raw_text = extract_pdf_text(file_bytes)

            if not raw_text.strip():
                json_response(self, 400, {"error": "No readable text was found in the uploaded file."})
                return

            title = os.path.splitext(os.path.basename(filename))[0]
            parsed = parse_questions(raw_text, title)
            parsed["sourceFile"] = filename
            json_response(self, 200, parsed)
        except ValueError as exc:
            json_response(self, 400, {"error": str(exc)})
        except RuntimeError as exc:
            json_response(self, 500, {"error": str(exc)})
        except Exception:
            json_response(self, 500, {"error": "The quiz file could not be imported. Check the format and answer key."})

    def do_GET(self):
        json_response(self, 200, {
            "ok": True,
            "message": "POST a DOCX or PDF quiz file to import questions."
        })
